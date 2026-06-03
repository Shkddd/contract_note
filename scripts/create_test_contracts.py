"""Generate two test contract documents for ContractReview testing.

Contract 1 — PDF: 采购合同（有违约风险条款）
  - match: 付款条款, 争议解决, 不可抗力
  - conflict: 违约责任仅5%违约金（KB要求不低于合同额）
  - missing: 保密条款缺失

Contract 2 — DOCX: 技术服务合同（含知识产权风险）
  - match: 保密义务, 知识产权归属, 数据保护
  - conflict: 竞业限制仅6个月（KB要求不少于12个月）
  - missing: 自动续期条款缺失（KB建议避免自动续期）
"""

import os
from pathlib import Path

OUTPUT_DIR = Path.home() / "contract-review" / "backend" / "data" / "uploads"

# ══════════════════════════════════════════════════════════════
# Contract 1: PDF — 采购合同
# ══════════════════════════════════════════════════════════════

CONTRACT_1_CLAUSES = [
    ("第一条 合同标的", "甲方（采购方）向乙方（供应商）采购以下货物：\n"
     "1. 服务器设备 20台，型号DELL PowerEdge R760\n"
     "2. 网络交换机 5台，型号Cisco Catalyst 9300\n"
     "合同总金额：人民币贰佰叁拾万元整（¥2,300,000.00）"),

    ("第二条 交付与验收", "乙方应于合同签订后30个工作日内完成交付。\n"
     "甲方应在收到货物后7个工作日内完成验收。\n"
     "验收不合格的，乙方应在15个工作日内完成更换。"),

    ("第三条 付款条款", "合同总价款采用分期付款方式：\n"
     "1. 合同签订后7日内，甲方向乙方支付合同总额的30%作为预付款；\n"
     "2. 货物全部交付并经甲方验收合格后15日内，支付合同总额的60%；\n"
     "3. 剩余10%作为质保金，质保期满后14日内支付。"),

    ("第四条 违约责任", "任何一方违反本合同约定，应向守约方支付合同总金额5%的违约金。\n"
     "违约金不足以弥补守约方实际损失的，守约方有权要求违约方赔偿实际损失。\n"
     "因乙方原因逾期交付的，每逾期一日按合同总金额的0.1%支付违约金。"),

    ("第五条 不可抗力", "因自然灾害、战争、政府行为等不可抗力因素导致合同无法履行的，\n"
     "受影响方应在不可抗力事件发生后7日内书面通知对方，并提供相关证明。\n"
     "不可抗力事件持续超过30日的，任何一方有权解除合同。"),

    ("第六条 争议解决", "因本合同引起的或与本合同有关的任何争议，双方应友好协商解决。\n"
     "协商不成的，提交甲方所在地有管辖权的人民法院诉讼解决。\n"
     "诉讼过程中，除争议事项外，双方应继续履行本合同其他条款。"),

    ("第七条 其他约定", "本合同一式两份，甲乙双方各执一份，具有同等法律效力。\n"
     "本合同自双方签字盖章之日起生效。\n"
     "合同附件为合同不可分割的组成部分，与合同具有同等法律效力。"),
]


def create_pdf():
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_font("Songti", "", "/System/Library/Fonts/Supplemental/Songti.ttc", uni=True)
    pdf.set_auto_page_break(auto=True, margin=20)

    # Title page
    pdf.add_page()
    pdf.set_font("Songti", "", 24)
    pdf.ln(60)
    pdf.cell(0, 15, "采  购  合  同", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(10)
    pdf.set_font("Songti", "", 14)
    pdf.cell(0, 10, "（标准设备采购）", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(30)
    pdf.set_font("Songti", "", 12)
    info_lines = [
        "合同编号：CG-2026-0088",
        "甲方：深圳恒通电子科技有限公司",
        "乙方：北京华信科技发展有限公司",
        "签订日期：2026年6月1日",
        "签订地点：深圳市南山区",
    ]
    for line in info_lines:
        pdf.cell(0, 8, line, new_x="LMARGIN", new_y="NEXT", align="C")

    # Clauses
    for title, content in CONTRACT_1_CLAUSES:
        pdf.add_page()
        pdf.set_font("Songti", "", 16)
        pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)
        pdf.set_font("Songti", "", 11)
        for line in content.split("\n"):
            # Handle indented bullet items
            indent = "    " if line.startswith(("1.", "2.", "3.")) else ""
            pdf.multi_cell(0, 7, indent + line)
            pdf.ln(2)

    outpath = OUTPUT_DIR / "采购合同_CG-2026-0088.pdf"
    pdf.output(str(outpath))
    print(f"PDF created: {outpath}")
    return outpath


# ══════════════════════════════════════════════════════════════
# Contract 2: DOCX — 技术服务合同
# ══════════════════════════════════════════════════════════════

CONTRACT_2_PARAGRAPHS = [
    ("技术服务合同", "heading1"),
    ("合同编号：JS-2026-0156", "normal"),
    ("甲方：上海前沿科技有限公司\n乙方：杭州智云信息技术有限公司\n签订日期：2026年6月1日", "normal"),
    ("", "normal"),

    ("第一条 服务内容", "heading2"),
    ("乙方为甲方提供企业级AI客服系统的技术开发服务，包括但不限于：\n"
     "1. 智能对话引擎开发与部署\n"
     "2. 知识库管理系统定制\n"
     "3. 第三方系统集成（CRM/ERP）\n"
     "4. 系统运维支持（12个月）", "body"),

    ("第二条 服务期限与费用", "heading2"),
    ("服务期限：自2026年6月15日至2027年6月14日，共计12个月。\n"
     "合同总金额：人民币壹佰伍拾万元整（¥1,500,000.00）\n"
     "费用支付：项目启动支付40%，中期验收支付30%，终验支付30%。", "body"),

    ("第三条 保密义务", "heading2"),
    ("双方应对在合作过程中获悉的对方商业秘密和机密信息承担保密义务。\n"
     "保密信息包括但不限于：技术资料、源代码、客户数据、商业计划等。\n"
     "保密期限自本合同终止之日起继续有效3年。\n"
     "未经对方书面同意，任何一方不得向第三方披露保密信息，法律法规另有规定的除外。\n"
     "违反保密义务的一方应赔偿对方因此遭受的全部直接和间接损失。", "body"),

    ("第四条 知识产权归属", "heading2"),
     ("乙方在履行本合同过程中独立开发完成的技术成果，其知识产权归属乙方所有。\n"
     "甲方在合同范围内享有该技术成果的永久、免费使用许可。\n"
     "双方共同开发完成的成果，知识产权由双方共同所有。\n"
     "甲方提供的原有知识产权仍归甲方所有。\n"
     "乙方保证其提供的技术成果不侵犯任何第三方的知识产权。", "body"),

    ("第五条 竞业限制", "heading2"),
    ("乙方及其项目核心人员在本合同履行期间及合同终止后6个月内，\n"
     "不得为甲方的竞争对手提供与本合同相同或类似的技术服务。\n"
     "竞业限制期间，甲方无需支付竞业限制补偿金。", "body"),

    ("第六条 数据保护", "heading2"),
    ("乙方在处理甲方数据时应遵守《个人信息保护法》和《数据安全法》的相关规定。\n"
     "乙方应采取必要的技术措施确保数据安全，防止数据泄露、篡改或丢失。\n"
     "乙方不得将甲方数据用于本合同约定之外的任何目的。\n"
     "服务期限届满后，乙方应按甲方要求删除或返还所有甲方数据。", "body"),

    ("第七条 违约责任", "heading2"),
    ("任何一方违反本合同约定，应向守约方支付合同总金额10%的违约金。\n"
     "因乙方原因导致项目延期交付的，每延期一日按合同总金额的0.2%支付违约金。\n"
     "任何一方累计违约超过30日的，守约方有权单方解除合同。", "body"),

    ("第八条 争议解决", "heading2"),
    ("因本合同引起的争议，双方应友好协商解决。协商不成的，\n"
     "任何一方可提交杭州仲裁委员会按照其仲裁规则进行仲裁。\n"
     "仲裁裁决是终局的，对双方均有约束力。", "body"),

    ("第九条 通知与送达", "heading2"),
    ("双方因履行本合同而相互发送的通知，应采用书面形式。\n"
     "通知可通过专人送达、挂号信、快递或电子邮件等方式发送。\n"
     "以下地址为双方的有效送达地址：\n"
     "甲方：上海市浦东新区张江高科技园区……\n"
     "乙方：杭州市余杭区未来科技城……", "body"),

    ("签署页", "heading2"),
    ("甲方（盖章）：上海前沿科技有限公司\n授权代表：张伟\n日期：2026年6月1日\n\n"
     "乙方（盖章）：杭州智云信息技术有限公司\n授权代表：李明\n日期：2026年6月1日", "body"),
]


def create_docx():
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Songti SC"
    font.size = Pt(11)

    for text, style_name in CONTRACT_2_PARAGRAPHS:
        if not text and style_name == "normal":
            continue
        if style_name == "heading1":
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style_name == "heading2":
            p = doc.add_heading(text, level=2)
        elif style_name == "body":
            for para_text in text.split("\n"):
                p = doc.add_paragraph(para_text, style="Normal")
                if para_text.startswith(("1.", "2.", "3.", "4.", "5.")):
                    p.paragraph_format.left_indent = Inches(0.3)
        else:
            p = doc.add_paragraph(text, style="Normal")

    outpath = OUTPUT_DIR / "技术服务合同_JS-2026-0156.docx"
    doc.save(str(outpath))
    print(f"DOCX created: {outpath}")
    return outpath


if __name__ == "__main__":
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    create_pdf()
    create_docx()
    print("\nDone! Both test contracts created.")

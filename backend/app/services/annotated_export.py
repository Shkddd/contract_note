"""Generate AI-annotated versions of reviewed contracts (PDF/DOCX).

Output uses native annotation format:
  - DOCX: Word comment bubbles (right sidebar, hover/click to view)
  - PDF:  PDF text annotations (comment icon, click for popup)
"""

import os
import re
import io
import zipfile
import datetime
from pathlib import Path
from typing import Optional, List
from datetime import datetime as dt
from xml.etree import ElementTree as ET

from .db import get_db

# ── Helpers ──────────────────────────────────────────────────────────

_RISK_COLORS = {"高": (200, 60, 60), "中": (200, 160, 40), "低": (60, 160, 60)}
_RISK_COLORS_HEX = {"高": "c83c3c", "中": "c8a028", "低": "3ca03c"}
_RISK_HEX_SHORT = {"高": "#E53935", "中": "#FB8C00", "低": "#43A047"}
_MATCH_LABELS = {
    "conflict": "[冲突]",
    "matched": "[匹配]",
    "missing": "[缺失]",
}

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS_CT = "http://schemas.openxmlformats.org/package/2006/content-types"


def _build_comment_text(ann: dict) -> str:
    lines = [
        f"条款 {ann['clause_index']} — {_MATCH_LABELS.get(ann['match_type'], ann['match_type'])} — {ann['risk_level']}风险",
        f"参考: {ann.get('kb_title', '—')}",
        "",
        ann.get("comment", ""),
    ]
    if ann.get("suggestion"):
        lines.extend(["", f"建议: {ann['suggestion']}"])
    return "\n".join(lines)


# ── Data ──────────────────────────────────────────────────────────────

def get_review_with_annotations(doc_id: int) -> Optional[dict]:
    """Fetch a completed review with all annotations for a document."""
    db = get_db()
    doc = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,)).fetchone()
    if not doc:
        return None
    doc = dict(doc)

    review = db.execute(
        "SELECT * FROM reviews WHERE doc_id = ? AND status = 'completed' ORDER BY id DESC LIMIT 1",
        (doc_id,),
    ).fetchone()
    if not review:
        return None
    review = dict(review)

    annotations = [
        dict(r) for r in db.execute(
            "SELECT a.*, c.content as clause_content, c.clause_index "
            "FROM annotations a "
            "JOIN clauses c ON c.id = a.clause_id "
            "WHERE c.doc_id = ? "
            "ORDER BY c.clause_index",
            (doc_id,),
        ).fetchall()
    ]

    clauses = [
        dict(r) for r in db.execute(
            "SELECT * FROM clauses WHERE doc_id = ? ORDER BY clause_index",
            (doc_id,),
        ).fetchall()
    ]

    return {
        "doc": doc,
        "review": review,
        "annotations": annotations,
        "clauses": clauses,
    }


# ── PDF Export ────────────────────────────────────────────────────────

def export_annotated_pdf(data: dict, output_path: str) -> str:
    """Generate PDF with native text annotations (comment popups)."""
    from fpdf import FPDF
    from fpdf.enums import AnnotationName

    pdf = FPDF()
    pdf.add_font("Songti", "", "/System/Library/Fonts/Supplemental/Songti.ttc")
    pdf.add_font("SongtiBold", "", "/System/Library/Fonts/Supplemental/Songti.ttc")
    pdf.set_auto_page_break(auto=True, margin=20)

    # ── Title Page ──
    pdf.add_page()
    pdf.set_font("Songti", "", 22)
    pdf.ln(50)
    pdf.cell(0, 14, "合同智能审核批注版", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(6)
    pdf.set_font("Songti", "", 12)
    pdf.cell(0, 8, data["doc"]["original_name"], new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(4)
    pdf.set_font("Songti", "", 10)

    info = data["review"]
    for line in [
        f"审核时间: {info['created_at']}",
        f"条款总数: {info['total_clauses']} 条",
        f"高风险: {info['high_risk']} · 中风险: {info['medium_risk']} · 低风险: {info['low_risk']}",
        f"匹配: {info['matched']} · 冲突: {info['conflicted']} · 缺失: {info['missing_info']}",
    ]:
        pdf.cell(0, 7, line, new_x="LMARGIN", new_y="NEXT", align="C")

    # Legend
    pdf.ln(6)
    pdf.set_font("Songti", "", 10)
    for level, hexcolor, label in [
        ("高", "E53935", "高风险"), ("中", "FB8C00", "中风险"), ("低", "43A047", "低风险"),
    ]:
        pdf.set_fill_color(int(hexcolor[:2], 16), int(hexcolor[2:4], 16), int(hexcolor[4:], 16))
        pdf.cell(8, 6, "", fill=True)
        pdf.cell(4)
        pdf.cell(0, 6, label, new_x="LMARGIN", new_y="NEXT", align="L")

    # ── Annotations Index ──
    pdf.add_page()
    pdf.set_font("Songti", "", 14)
    pdf.cell(0, 10, "批注意见总览", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    for ann in data["annotations"]:
        risk = ann["risk_level"]
        r, g_, b = (_RISK_COLORS_HEX[risk][i:i+2] for i in (0, 2, 4))
        pdf.set_fill_color(int(r, 16), int(g_, 16), int(b, 16))
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Songti", "", 9)
        tag = f"[{risk}] {_MATCH_LABELS.get(ann['match_type'], ann['match_type'])}"
        tw = pdf.get_string_width(tag) + 4
        pdf.cell(tw, 6, tag, fill=True)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(2)
        pdf.set_font("Songti", "", 9)
        pdf.cell(0, 6, f"条款{ann['clause_index']} — {ann.get('kb_title', '')}",
                new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Songti", "", 8)
        pdf.set_text_color(100, 100, 100)
        pdf.multi_cell(0, 5, ann["comment"][:120] + ("..." if len(ann["comment"]) > 120 else ""),
                       new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)

    # ── Clause-by-Clause with Text Annotations ──
    pdf.add_page()
    pdf.set_font("Songti", "", 14)
    pdf.cell(0, 10, "逐条审核详情", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    clause_anns = {}
    for ann in data["annotations"]:
        clause_anns.setdefault(ann["clause_index"], []).append(ann)

    for clause in data["clauses"]:
        ci = clause["clause_index"]
        anns = clause_anns.get(ci, [])

        if pdf.get_y() > 220:
            pdf.add_page()

        # ── Clause header bar ──
        y0 = pdf.get_y()
        has_risk = anns and anns[0]["risk_level"]
        bar_color = {"高": 200, "中": 180, "低": 140}.get(has_risk, 220)
        pdf.set_fill_color(bar_color, bar_color, 240)
        pdf.set_draw_color(180, 180, 220)
        pdf.set_line_width(0.5)
        pdf.rect(10, y0, 190, 10)
        pdf.set_xy(14, y0 + 1)
        pdf.set_font("Songti", "", 11)
        pdf.set_text_color(50, 50, 80)
        pdf.cell(0, 8, f"条款 {ci}",
                new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        # ── Clause text ──
        pdf.set_text_color(30, 30, 30)
        pdf.set_font("Songti", "", 9)
        clause_text = clause["content"].strip()
        for line in clause_text.split("\n"):
            if pdf.get_y() > 265:
                pdf.add_page()
            pdf.multi_cell(0, 5, line, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        # ── Insert comment annotation icon ──
        if anns:
            for ann in anns:
                if pdf.get_y() > 260:
                    pdf.add_page()

                # Small colored indicator
                risk = ann["risk_level"]
                r, g_, b = (_RISK_COLORS_HEX[risk][i:i+2] for i in (0, 2, 4))
                pdf.set_fill_color(int(r, 16), int(g_, 16), int(b, 16))
                tag = f"  {_MATCH_LABELS.get(ann['match_type'], ann['match_type'])} | {risk}风险  "
                tw = pdf.get_string_width(tag) + 4
                pdf.cell(tw, 6, tag, fill=True)

                # Place a text annotation (comment bubble) right after the tag
                # x, y are in mm from top-left; text_annotation takes bottom-left in user coords
                ann_x = pdf.get_x()
                ann_y = pdf.get_y()
                pdf.set_text_color(80, 80, 80)
                pdf.set_font("Songti", "", 8)
                pdf.cell(0, 6, f"  {ann.get('kb_title', '')}",
                        new_x="LMARGIN", new_y="NEXT")

                # Place a text_annotation (comment bubble icon) at this line
                comment_text = _build_comment_text(ann)
                pdf.text_annotation(
                    x=ann_x - tw - 5,  # position the icon next to the tag
                    y=ann_y,
                    w=4,
                    h=4,
                    text=comment_text,
                    title=f"条款{ci} — {ann['risk_level']}风险",
                    name="Comment",
                )

                # Also highlight the clause heading with a highlight annotation
                # (we'll skip highlight to keep it simple)

                pdf.ln(2)
        else:
            pdf.set_text_color(100, 180, 100)
            pdf.set_font("Songti", "", 9)
            pdf.cell(0, 6, "  ✓ 此条款无异常", new_x="LMARGIN", new_y="NEXT")

        pdf.ln(3)

    pdf.output(str(output_path))
    return output_path


# ── DOCX Export ──────────────────────────────────────────────────────


def _risk_highlight_color(risk_level: str):
    """Map risk level to a python-docx highlight color constant."""
    from docx.enum.text import WD_COLOR_INDEX
    return {"高": WD_COLOR_INDEX.RED, "中": WD_COLOR_INDEX.YELLOW,
            "低": WD_COLOR_INDEX.BRIGHT_GREEN}.get(risk_level, None)


def export_annotated_docx(data: dict, output_path: str) -> str:
    """Generate DOCX with native Word comment bubbles (+ highlights)."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Songti SC"
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.2

    # ── Title ──
    p = doc.add_heading("合同智能审核批注版", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(data["doc"]["original_name"])
    run.font.size = Pt(12)

    rev = data["review"]
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        f"审核时间: {rev['created_at']}  "
        f"高: {rev['high_risk']}  "
        f"中: {rev['medium_risk']}  "
        f"低: {rev['low_risk']}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph("─" * 60)

    # ── Summary ──
    doc.add_heading("审核概览", level=2)
    sum_p = doc.add_paragraph()
    run = sum_p.add_run(
        f"共 {rev['total_clauses']} 条条款 · "
        f"匹配 {rev['matched']} · "
        f"冲突 {rev['conflicted']} · "
        f"缺失 {rev['missing_info']}"
    )

    doc.add_heading("条款审核详情", level=2)

    # Build clause → annotations map
    clause_anns = {}
    for ann in data["annotations"]:
        clause_anns.setdefault(ann["clause_index"], []).append(ann)

    # Collect paragraphs that need comments
    comment_targets: List[tuple] = []

    for clause in data["clauses"]:
        ci = clause["clause_index"]
        anns = clause_anns.get(ci, [])

        # ── Clause heading ──
        p = doc.add_paragraph()
        run = p.add_run(f"条款 {ci}  ")
        run.bold = True
        run.font.size = Pt(11)

        if anns:
            # Add risk + match tag next to heading
            risk = anns[0]["risk_level"]
            hex_c = _RISK_HEX_SHORT[risk]
            tag_text = f"[{risk}] {_MATCH_LABELS.get(anns[0]['match_type'], anns[0]['match_type'])}"
            run_tag = p.add_run(f"  {tag_text}")
            run_tag.font.size = Pt(9)
            run_tag.bold = True
            rc = _RISK_COLORS[risk]
            run_tag.font.color.rgb = RGBColor(*rc)

            # Build multi-line comment text for the Word comment bubble
            parts = []
            for ann in anns:
                parts.append(f"条款 {ann['clause_index']} — "
                             f"{_MATCH_LABELS.get(ann['match_type'], ann['match_type'])} — "
                             f"{ann['risk_level']}风险")
                parts.append(f"参考: {ann.get('kb_title', '—')}")
                parts.append("")
                parts.append(ann.get("comment", ""))
                if ann.get("suggestion"):
                    parts.append("")
                    parts.append(f"建议: {ann['suggestion']}")
                parts.append("─" * 30)
            comment_text = "\n".join(parts).rstrip("─\n ")

            # Schedule comment for this paragraph
            comment_targets.append((p, comment_text, "ContractReview AI"))

        # ── Clause content ──
        clause_text = clause["content"].strip()
        for line in clause_text.split("\n"):
            cp = doc.add_paragraph(line, style="Normal")
            cp.paragraph_format.left_indent = Cm(0.5)
            if anns:
                # Highlight the first line to show it has annotations
                for run in cp.runs:
                    run.font.highlight_color = _risk_highlight_color(anns[0]["risk_level"])

        # Separator
        if not anns:
            p_ok = doc.add_paragraph()
            run_ok = p_ok.add_run("✓ 此条款无异常")
            run_ok.font.color.rgb = RGBColor(67, 160, 71)
            run_ok.font.size = Pt(9)

        doc.add_paragraph("─" * 40)

    # ── Attach native Word comments (post-save ZIP injection) ──
    if comment_targets:
        doc.save(str(output_path))
        _inject_comments_post_save(str(output_path), comment_targets)
    else:
        doc.save(str(output_path))
    return output_path


def _risk_highlight_color(risk_level: str):
    """Map risk level to a python-docx highlight color constant."""
    from docx.enum.text import WD_COLOR_INDEX
    return {"高": WD_COLOR_INDEX.RED, "中": WD_COLOR_INDEX.YELLOW,
            "低": WD_COLOR_INDEX.BRIGHT_GREEN}.get(risk_level, None)


def _inject_comments_post_save(docx_path: str, comment_targets: List[tuple]):
    """Inject native Word comments into a saved DOCX by manipulating the ZIP.

    Each comment_targets item: (paragraph, comment_text, author_name).
    The paragraph reference is used only for its XML element position.
    """
    import zipfile
    import io
    from lxml import etree as ETX

    W = f"{{{NS_W}}}"
    now_str = dt.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")

    # ── Read the saved docx into memory ──
    with open(docx_path, "rb") as f:
        original = f.read()

    # ── Build comments.xml ──
    comments_el = ETX.Element(f"{W}comments",
                              nsmap={"w": NS_W, "r": NS_R})
    for i, (paragraph, text, author) in enumerate(comment_targets):
        cid = i + 1
        comment = ETX.SubElement(comments_el, f"{W}comment")
        comment.set(f"{W}id", str(cid))
        comment.set(f"{W}author", author)
        comment.set(f"{W}date", now_str)

        # Para 1 — annotationRef (user icon in comment bubble)
        cp = ETX.SubElement(comment, f"{W}p")
        cr = ETX.SubElement(cp, f"{W}r")
        crpr = ETX.SubElement(cr, f"{W}rPr")
        crs = ETX.SubElement(crpr, f"{W}rStyle")
        crs.set(f"{W}val", "CommentReference")
        ETX.SubElement(cr, f"{W}annotationRef")

        # Para 2 — actual text
        for line in text.split("\n"):
            tp = ETX.SubElement(comment, f"{W}p")
            tr = ETX.SubElement(tp, f"{W}r")
            trpr = ETX.SubElement(tr, f"{W}rPr")
            trsz = ETX.SubElement(trpr, f"{W}sz")
            trsz.set(f"{W}val", "20")
            tt = ETX.SubElement(tr, f"{W}t")
            tt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            tt.text = line if line else " "

    comments_xml = ETX.tostring(comments_el, xml_declaration=True,
                                encoding="UTF-8", standalone=True)

    # ── Now process the ZIP ──
    buf_in = io.BytesIO(original)
    buf_out = io.BytesIO()

    with zipfile.ZipFile(buf_in, "r") as zin:
        with zipfile.ZipFile(buf_out, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)

                if item.filename == "word/document.xml":
                    # Patch document.xml: add comment markers
                    doc_root = ETX.fromstring(data)

                    # Register namespaces for serialization
                    ETX.register_namespace("w", NS_W)
                    ETX.register_namespace("r", NS_R)

                    for i, (paragraph, text, author) in enumerate(comment_targets):
                        cid = i + 1
                        p_elem = paragraph._element

                        # Find this paragraph in the re-parsed tree
                        # Match by unique content: first text run
                        p_id = None
                        for p_node in doc_root.iter(f"{W}p"):
                            first_text = p_node.find(f".//{W}t")
                            if first_text is not None and first_text.text:
                                if first_text.text.strip():
                                    # Match against original paragraph's first text
                                    orig_first = p_elem.find(f".//{W}t")
                                    if orig_first is not None and orig_first.text:
                                        if first_text.text.strip() == orig_first.text.strip():
                                            p_id = p_node
                                            break

                        if p_id is None:
                            continue

                        # commentRangeStart before paragraph
                        cr_start = ETX.Element(f"{W}commentRangeStart")
                        cr_start.set(f"{W}id", str(cid))
                        parent = p_id.getparent()
                        parent.insert(list(parent).index(p_id), cr_start)

                        # commentReference inside paragraph (after first run)
                        first_r = p_id.find(f"{W}r")
                        if first_r is not None:
                            ref_run = ETX.Element(f"{W}r")
                            ref_rpr = ETX.SubElement(ref_run, f"{W}rPr")
                            ref_rs = ETX.SubElement(ref_rpr, f"{W}rStyle")
                            ref_rs.set(f"{W}val", "CommentReference")
                            ETX.SubElement(ref_run, f"{W}commentReference",
                                          {f"{W}id": str(cid)})
                            p_id.insert(list(p_id).index(first_r), ref_run)

                        # commentRangeEnd after paragraph
                        cr_end = ETX.Element(f"{W}commentRangeEnd")
                        cr_end.set(f"{W}id", str(cid))
                        next_el = p_id.getnext()
                        if next_el is not None:
                            next_el.addprevious(cr_end)
                        else:
                            parent.append(cr_end)

                        # End-of-para commentReference
                        tail_ref = ETX.SubElement(p_id, f"{W}r")
                        tail_rpr = ETX.SubElement(tail_ref, f"{W}rPr")
                        tail_rs = ETX.SubElement(tail_rpr, f"{W}rStyle")
                        tail_rs.set(f"{W}val", "CommentReference")
                        ETX.SubElement(tail_ref, f"{W}commentReference",
                                      {f"{W}id": str(cid)})

                    patched_doc = ETX.tostring(doc_root, xml_declaration=True,
                                              encoding="UTF-8", standalone=True)
                    zout.writestr(item.filename, patched_doc)

                elif item.filename == "[Content_Types].xml":
                    # Add comments content type
                    ct_root = ETX.fromstring(data)
                    ct_ns = "{http://schemas.openxmlformats.org/package/2006/content-types}"
                    has_ct = False
                    for override in ct_root.findall(f"{ct_ns}Override"):
                        if override.get("PartName") == "/word/comments.xml":
                            has_ct = True
                            break
                    if not has_ct:
                        override = ETX.SubElement(ct_root, f"{ct_ns}Override")
                        override.set("PartName", "/word/comments.xml")
                        override.set("ContentType",
                                     "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml")
                    patched_ct = ETX.tostring(ct_root, xml_declaration=True,
                                             encoding="UTF-8", standalone=True)
                    zout.writestr(item.filename, patched_ct)

                elif item.filename == "word/_rels/document.xml.rels":
                    # Add comments relationship
                    rels_root = ETX.fromstring(data)
                    rels_ns = "{http://schemas.openxmlformats.org/package/2006/relationships}"
                    has_rel = False
                    for rel in rels_root.findall(f"{rels_ns}Relationship"):
                        if rel.get("Target") == "comments.xml":
                            has_rel = True
                            break
                    if not has_rel:
                        rel_elem = ETX.SubElement(rels_root, f"{rels_ns}Relationship")
                        rel_elem.set("Id", "rComments1")
                        rel_elem.set("Type",
                                     "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments")
                        rel_elem.set("Target", "comments.xml")
                    patched_rels = ETX.tostring(rels_root, xml_declaration=True,
                                               encoding="UTF-8", standalone=True)
                    zout.writestr(item.filename, patched_rels)

                else:
                    zout.writestr(item, data)

            # Write comments.xml
            zout.writestr("word/comments.xml", comments_xml)

    # ── Write back to the same path ──
    with open(docx_path, "wb") as f:
        f.write(buf_out.getvalue())


# ── Public Entry Point ────────────────────────────────────────────────

def export_annotated(doc_id: int, output_dir: str) -> dict:
    """Generate annotated document. Returns {filepath, filename, mime}."""
    data = get_review_with_annotations(doc_id)
    if not data:
        return {"error": "未找到审核结果"}

    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    ext = data["doc"]["file_type"]
    orig_name = Path(data["doc"]["original_name"]).stem
    ts = dt.now().strftime("%Y%m%d_%H%M%S")

    if ext == "pdf":
        out_path = str(out_dir / f"{orig_name}_批注版_{ts}.pdf")
        export_annotated_pdf(data, out_path)
        return {
            "filepath": out_path,
            "filename": f"{orig_name}_批注版.pdf",
            "mime": "application/pdf",
        }
    elif ext == "docx":
        out_path = str(out_dir / f"{orig_name}_批注版_{ts}.docx")
        export_annotated_docx(data, out_path)
        return {
            "filepath": out_path,
            "filename": f"{orig_name}_批注版.docx",
            "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
    else:
        return {"error": f"不支持的文件类型: {ext}"}
